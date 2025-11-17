"""
File Upload & Parsing Manager for iTechSmart Ninja
Handles multi-format file uploads and content extraction
"""

import logging
import uuid
import mimetypes
from typing import Dict, List, Optional, Any, BinaryIO
from datetime import datetime
from dataclasses import dataclass, asdict
from enum import Enum
import os
import hashlib

logger = logging.getLogger(__name__)


class FileType(str, Enum):
    """Supported file types"""

    PDF = "pdf"
    DOCX = "docx"
    DOC = "doc"
    XLSX = "xlsx"
    XLS = "xls"
    CSV = "csv"
    TXT = "txt"
    JSON = "json"
    XML = "xml"
    HTML = "html"
    MARKDOWN = "markdown"
    RTF = "rtf"
    IMAGE = "image"
    AUDIO = "audio"
    VIDEO = "video"
    ARCHIVE = "archive"
    CODE = "code"
    UNKNOWN = "unknown"


class ParsingStatus(str, Enum):
    """File parsing status"""

    PENDING = "pending"
    PARSING = "parsing"
    COMPLETED = "completed"
    FAILED = "failed"


@dataclass
class FileMetadata:
    """File metadata information"""

    file_id: str
    filename: str
    file_type: FileType
    mime_type: str
    size_bytes: int
    hash_md5: str
    hash_sha256: str
    uploaded_at: datetime
    uploaded_by: str

    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary"""
        return {
            "file_id": self.file_id,
            "filename": self.filename,
            "file_type": self.file_type.value,
            "mime_type": self.mime_type,
            "size_bytes": self.size_bytes,
            "hash_md5": self.hash_md5,
            "hash_sha256": self.hash_sha256,
            "uploaded_at": self.uploaded_at.isoformat(),
            "uploaded_by": self.uploaded_by,
        }


@dataclass
class ParsedContent:
    """Parsed file content"""

    file_id: str
    status: ParsingStatus
    content_type: str
    text_content: Optional[str]
    structured_data: Optional[Dict[str, Any]]
    metadata: Dict[str, Any]
    page_count: Optional[int]
    word_count: Optional[int]
    parsed_at: Optional[datetime]
    error_message: Optional[str]

    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary"""
        return {
            "file_id": self.file_id,
            "status": self.status.value,
            "content_type": self.content_type,
            "text_content": self.text_content,
            "structured_data": self.structured_data,
            "metadata": self.metadata,
            "page_count": self.page_count,
            "word_count": self.word_count,
            "parsed_at": self.parsed_at.isoformat() if self.parsed_at else None,
            "error_message": self.error_message,
        }


class FileParser:
    """Handles file parsing for different formats"""

    @staticmethod
    def parse_text(content: bytes) -> str:
        """Parse plain text file"""
        try:
            return content.decode("utf-8")
        except UnicodeDecodeError:
            return content.decode("latin-1")

    @staticmethod
    def parse_json(content: bytes) -> Dict[str, Any]:
        """Parse JSON file"""
        import json

        return json.loads(content.decode("utf-8"))

    @staticmethod
    def parse_csv(content: bytes) -> List[Dict[str, Any]]:
        """Parse CSV file"""
        import csv
        import io

        text = content.decode("utf-8")
        reader = csv.DictReader(io.StringIO(text))
        return list(reader)

    @staticmethod
    def parse_xml(content: bytes) -> Dict[str, Any]:
        """Parse XML file"""
        import xml.etree.ElementTree as ET

        root = ET.fromstring(content)

        def element_to_dict(element):
            result = {}
            if element.text and element.text.strip():
                result["_text"] = element.text.strip()

            for child in element:
                child_data = element_to_dict(child)
                if child.tag in result:
                    if not isinstance(result[child.tag], list):
                        result[child.tag] = [result[child.tag]]
                    result[child.tag].append(child_data)
                else:
                    result[child.tag] = child_data

            result.update(element.attrib)
            return result

        return {root.tag: element_to_dict(root)}

    @staticmethod
    def parse_pdf(file_path: str) -> Dict[str, Any]:
        """Parse PDF file"""
        try:
            import PyPDF2

            with open(file_path, "rb") as file:
                reader = PyPDF2.PdfReader(file)

                text_content = []
                for page in reader.pages:
                    text_content.append(page.extract_text())

                return {
                    "text": "\n\n".join(text_content),
                    "page_count": len(reader.pages),
                    "metadata": reader.metadata,
                }
        except ImportError:
            logger.warning("PyPDF2 not installed, using basic extraction")
            return {"text": "PDF parsing requires PyPDF2", "page_count": 0}

    @staticmethod
    def parse_docx(file_path: str) -> Dict[str, Any]:
        """Parse DOCX file"""
        try:
            from docx import Document

            doc = Document(file_path)

            text_content = []
            for paragraph in doc.paragraphs:
                text_content.append(paragraph.text)

            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)

            return {
                "text": "\n".join(text_content),
                "tables": tables,
                "paragraph_count": len(doc.paragraphs),
            }
        except ImportError:
            logger.warning("python-docx not installed")
            return {"text": "DOCX parsing requires python-docx"}

    @staticmethod
    def parse_xlsx(file_path: str) -> Dict[str, Any]:
        """Parse XLSX file"""
        try:
            import openpyxl

            workbook = openpyxl.load_workbook(file_path)

            sheets = {}
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]

                data = []
                for row in sheet.iter_rows(values_only=True):
                    data.append(list(row))

                sheets[sheet_name] = data

            return {
                "sheets": sheets,
                "sheet_names": workbook.sheetnames,
                "sheet_count": len(workbook.sheetnames),
            }
        except ImportError:
            logger.warning("openpyxl not installed")
            return {"text": "XLSX parsing requires openpyxl"}

    @staticmethod
    def parse_markdown(content: bytes) -> Dict[str, Any]:
        """Parse Markdown file"""
        text = content.decode("utf-8")

        # Extract headers
        headers = []
        for line in text.split("\n"):
            if line.startswith("#"):
                level = len(line) - len(line.lstrip("#"))
                title = line.lstrip("#").strip()
                headers.append({"level": level, "title": title})

        return {"text": text, "headers": headers}

    @staticmethod
    def parse_html(content: bytes) -> Dict[str, Any]:
        """Parse HTML file"""
        try:
            from bs4 import BeautifulSoup

            soup = BeautifulSoup(content, "html.parser")

            # Extract text
            text = soup.get_text()

            # Extract links
            links = [a.get("href") for a in soup.find_all("a") if a.get("href")]

            # Extract images
            images = [img.get("src") for img in soup.find_all("img") if img.get("src")]

            return {
                "text": text,
                "links": links,
                "images": images,
                "title": soup.title.string if soup.title else None,
            }
        except ImportError:
            logger.warning("beautifulsoup4 not installed")
            return {"text": content.decode("utf-8")}


class FileManager:
    """Manages file uploads and parsing"""

    def __init__(self, storage_path: str = "/workspace/uploads"):
        """Initialize file manager"""
        self.storage_path = storage_path
        self.files: Dict[str, FileMetadata] = {}
        self.parsed_content: Dict[str, ParsedContent] = {}

        # Create storage directory
        os.makedirs(storage_path, exist_ok=True)

        logger.info(f"FileManager initialized with storage path: {storage_path}")

    def _detect_file_type(self, filename: str, mime_type: str) -> FileType:
        """Detect file type from filename and mime type"""
        ext = os.path.splitext(filename)[1].lower()

        type_map = {
            ".pdf": FileType.PDF,
            ".docx": FileType.DOCX,
            ".doc": FileType.DOC,
            ".xlsx": FileType.XLSX,
            ".xls": FileType.XLS,
            ".csv": FileType.CSV,
            ".txt": FileType.TXT,
            ".json": FileType.JSON,
            ".xml": FileType.XML,
            ".html": FileType.HTML,
            ".htm": FileType.HTML,
            ".md": FileType.MARKDOWN,
            ".rtf": FileType.RTF,
            ".jpg": FileType.IMAGE,
            ".jpeg": FileType.IMAGE,
            ".png": FileType.IMAGE,
            ".gif": FileType.IMAGE,
            ".mp3": FileType.AUDIO,
            ".wav": FileType.AUDIO,
            ".mp4": FileType.VIDEO,
            ".avi": FileType.VIDEO,
            ".zip": FileType.ARCHIVE,
            ".tar": FileType.ARCHIVE,
            ".gz": FileType.ARCHIVE,
            ".py": FileType.CODE,
            ".js": FileType.CODE,
            ".java": FileType.CODE,
            ".cpp": FileType.CODE,
        }

        return type_map.get(ext, FileType.UNKNOWN)

    def _calculate_hashes(self, content: bytes) -> tuple:
        """Calculate MD5 and SHA256 hashes"""
        md5 = hashlib.md5(content).hexdigest()
        sha256 = hashlib.sha256(content).hexdigest()
        return md5, sha256

    async def upload_file(
        self,
        filename: str,
        content: bytes,
        user_id: str,
        mime_type: Optional[str] = None,
    ) -> FileMetadata:
        """
        Upload a file

        Args:
            filename: Original filename
            content: File content
            user_id: User ID
            mime_type: Optional MIME type

        Returns:
            FileMetadata object
        """
        file_id = str(uuid.uuid4())

        # Detect MIME type if not provided
        if mime_type is None:
            mime_type, _ = mimetypes.guess_type(filename)
            if mime_type is None:
                mime_type = "application/octet-stream"

        # Detect file type
        file_type = self._detect_file_type(filename, mime_type)

        # Calculate hashes
        md5, sha256 = self._calculate_hashes(content)

        # Save file
        file_path = os.path.join(self.storage_path, f"{file_id}_{filename}")
        with open(file_path, "wb") as f:
            f.write(content)

        # Create metadata
        metadata = FileMetadata(
            file_id=file_id,
            filename=filename,
            file_type=file_type,
            mime_type=mime_type,
            size_bytes=len(content),
            hash_md5=md5,
            hash_sha256=sha256,
            uploaded_at=datetime.now(),
            uploaded_by=user_id,
        )

        self.files[file_id] = metadata

        logger.info(f"File {filename} uploaded with ID {file_id}")
        return metadata

    async def parse_file(self, file_id: str) -> ParsedContent:
        """
        Parse a file and extract content

        Args:
            file_id: File ID

        Returns:
            ParsedContent object
        """
        metadata = self.files.get(file_id)
        if not metadata:
            raise ValueError(f"File {file_id} not found")

        file_path = os.path.join(self.storage_path, f"{file_id}_{metadata.filename}")

        parsed = ParsedContent(
            file_id=file_id,
            status=ParsingStatus.PARSING,
            content_type=metadata.file_type.value,
            text_content=None,
            structured_data=None,
            metadata={},
            page_count=None,
            word_count=None,
            parsed_at=None,
            error_message=None,
        )

        try:
            # Read file content
            with open(file_path, "rb") as f:
                content = f.read()

            parser = FileParser()

            # Parse based on file type
            if metadata.file_type == FileType.TXT:
                text = parser.parse_text(content)
                parsed.text_content = text
                parsed.word_count = len(text.split())

            elif metadata.file_type == FileType.JSON:
                data = parser.parse_json(content)
                parsed.structured_data = data
                parsed.text_content = str(data)

            elif metadata.file_type == FileType.CSV:
                data = parser.parse_csv(content)
                parsed.structured_data = {"rows": data}
                parsed.text_content = str(data)

            elif metadata.file_type == FileType.XML:
                data = parser.parse_xml(content)
                parsed.structured_data = data
                parsed.text_content = str(data)

            elif metadata.file_type == FileType.PDF:
                data = parser.parse_pdf(file_path)
                parsed.text_content = data.get("text")
                parsed.page_count = data.get("page_count")
                parsed.metadata = data.get("metadata", {})
                if parsed.text_content:
                    parsed.word_count = len(parsed.text_content.split())

            elif metadata.file_type == FileType.DOCX:
                data = parser.parse_docx(file_path)
                parsed.text_content = data.get("text")
                parsed.structured_data = {"tables": data.get("tables", [])}
                if parsed.text_content:
                    parsed.word_count = len(parsed.text_content.split())

            elif metadata.file_type == FileType.XLSX:
                data = parser.parse_xlsx(file_path)
                parsed.structured_data = data
                parsed.text_content = str(data)

            elif metadata.file_type == FileType.MARKDOWN:
                data = parser.parse_markdown(content)
                parsed.text_content = data.get("text")
                parsed.structured_data = {"headers": data.get("headers", [])}
                if parsed.text_content:
                    parsed.word_count = len(parsed.text_content.split())

            elif metadata.file_type == FileType.HTML:
                data = parser.parse_html(content)
                parsed.text_content = data.get("text")
                parsed.structured_data = {
                    "links": data.get("links", []),
                    "images": data.get("images", []),
                }
                parsed.metadata = {"title": data.get("title")}
                if parsed.text_content:
                    parsed.word_count = len(parsed.text_content.split())

            else:
                parsed.text_content = (
                    f"Parsing not supported for {metadata.file_type.value}"
                )

            parsed.status = ParsingStatus.COMPLETED
            parsed.parsed_at = datetime.now()

        except Exception as e:
            logger.error(f"Error parsing file {file_id}: {e}")
            parsed.status = ParsingStatus.FAILED
            parsed.error_message = str(e)

        self.parsed_content[file_id] = parsed
        return parsed

    async def get_file_metadata(self, file_id: str) -> Optional[FileMetadata]:
        """Get file metadata"""
        return self.files.get(file_id)

    async def get_parsed_content(self, file_id: str) -> Optional[ParsedContent]:
        """Get parsed content"""
        return self.parsed_content.get(file_id)

    async def delete_file(self, file_id: str) -> bool:
        """Delete a file"""
        metadata = self.files.get(file_id)
        if not metadata:
            raise ValueError(f"File {file_id} not found")

        file_path = os.path.join(self.storage_path, f"{file_id}_{metadata.filename}")

        try:
            if os.path.exists(file_path):
                os.remove(file_path)

            del self.files[file_id]

            if file_id in self.parsed_content:
                del self.parsed_content[file_id]

            logger.info(f"File {file_id} deleted")
            return True
        except Exception as e:
            logger.error(f"Error deleting file {file_id}: {e}")
            raise

    async def list_files(
        self, user_id: Optional[str] = None, file_type: Optional[FileType] = None
    ) -> List[FileMetadata]:
        """List files with optional filtering"""
        files = list(self.files.values())

        if user_id:
            files = [f for f in files if f.uploaded_by == user_id]

        if file_type:
            files = [f for f in files if f.file_type == file_type]

        return files


# Global file manager instance
_file_manager: Optional[FileManager] = None


def get_file_manager() -> FileManager:
    """Get or create global file manager instance"""
    global _file_manager
    if _file_manager is None:
        _file_manager = FileManager()
    return _file_manager
